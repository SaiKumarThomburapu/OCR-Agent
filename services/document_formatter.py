"""Professional Word Document Formatter - Bakeer Academy DOCX (PRODUCTION)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Any
import os

class DocxFormatter:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup Bakeer Academy professional styles"""
        # Set default font
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(12)
    
    def create_bakeer_docx(self, questions: List[Dict[str, Any]], output_path: str):
        """Create professional DOCX in Bakeer Academy format"""
        self.doc = Document()
        
        # Page 1: Header with branding
        self._create_header_page()
        
        # Pages: Questions
        for i, q in enumerate(questions, 1):
            self._create_question_page(q, i, len(questions))
        
        # Save
        self.doc.save(output_path)
        print(f"✅ Created {output_path} ({len(questions)} questions)")
        return output_path
    
    def _create_header_page(self):
        """Create professional header page"""
        # Add spacing
        self.doc.add_paragraph()
        
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("BAKEER ACADEMY")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)  # Dark blue
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Mathematics Questions & Answers")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 102, 204)
        
        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Info box
        info = self.doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run("Prepared for: Bakeer Academy GAT Program\nPhone: 0552420800")
        run.font.size = Pt(12)
        run.font.italic = True
    
    def _create_question_page(self, question: Dict[str, Any], q_num: int, total: int):
        """Create professional question page - BAKEER FORMAT"""
        
        # Add page break
        self.doc.add_page_break()
        
        # Question number in header
        header = self.doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run(f"Question {q_num}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Question text box
        question_para = self.doc.add_paragraph()
        question_para.paragraph_format.left_indent = Inches(0.5)
        question_para.paragraph_format.space_before = Pt(12)
        question_para.paragraph_format.space_after = Pt(18)
        
        # Question text
        run = question_para.add_run(question.get("question_text", ""))
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Options table (Bakeer style)
        self._add_options_table(question)
        
        # Answer (if available)
        if question.get("correct_answer"):
            answer_para = self.doc.add_paragraph()
            answer_para.paragraph_format.left_indent = Inches(0.5)
            answer_para.paragraph_format.space_before = Pt(12)
            
            run = answer_para.add_run(f"Answer: {question['correct_answer']}")
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)
    
    def _add_options_table(self, question: Dict[str, Any]):
        """Add options as professional table (Bakeer style)"""
        options = question.get("options", {})
        
        # Create 2x2 table for 4 options
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Set table width
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(2.5)
        
        option_list = [
            (options.get("A", ""), "A"),
            (options.get("B", ""), "B"),
            (options.get("C", ""), "C"),
            (options.get("D", ""), "D"),
        ]
        
        # Fill table
        cells = [cell for row in table.rows for cell in row.cells]
        for i, (cell, (option_text, label)) in enumerate(zip(cells, option_list)):
            # Clear cell
            cell.text = ""
            
            # Add option
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run(f"{label}) {option_text}")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
            
            # Set cell background
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F0F0')
            cell._element.get_or_add_tcPr().append(shading_elm)

def test_docx_formatter():
    """Test DOCX formatter"""
    formatter = DocxFormatter()
    sample_questions = [
        {
            "question_number": 1,
            "question_text": "Find the value (1/32 ÷ 1/4) × (1/16 ÷ 1/8) =",
            "options": {"A": "1/16", "B": "1/8", "C": "16/1", "D": "4/16"},
            "correct_answer": "A",
            "confidence": 0.95
        }
    ]
    output = formatter.create_bakeer_docx(sample_questions, "test_bakeer.docx")
    print(f"✅ Test DOCX created: {output}")

